"""별첨 DOCX 생성 → 바탕화면. (AI 활용 증빙 / 시제품 포트폴리오)

외부 평가 반영: AI 학습도구=XGBoost·분석도구=SHAP·Claude는 개발보조로 구분,
[직접 추가] 미완성 문구 제거, 커밋 수 표현 통일, 시제품에 시뮬 실제 수치.
Claude Code 대화 캡처는 선택(있으면 보강).
실행: py -3.12 backend/scripts/build_attachments_docx.py
"""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1].parent
EV = ROOT / "docs" / "submission" / "evidence"
DESK = Path(os.path.expanduser("~")) / "Desktop"


def set_kfont(doc, font="맑은 고딕"):
    for name in ["Normal", "Title", "Heading 1", "Heading 2", "Light Grid Accent 1"]:
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        st.font.name = font
        rpr = st.element.get_or_add_rPr()
        rf = rpr.get_or_add_rFonts()
        rf.set(qn("w:ascii"), font)
        rf.set(qn("w:hAnsi"), font)
        rf.set(qn("w:eastAsia"), font)


def caption(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def pic(doc, name, width=6.2):
    f = EV / name
    if f.exists():
        doc.add_picture(str(f), width=Inches(width))


# ───────────── 1) AI 활용 증빙 ─────────────
d = Document()
set_kfont(d)
d.add_heading("AI 활용 증빙 — BlindZone", 0)
d.add_paragraph("2026 국토교통 데이터 활용 경진대회 / 가점 신청 증빙. 부여 여부는 심사위원단 판단.")
d.add_paragraph("라이브: https://blindzone-brown.vercel.app  ·  코드: https://github.com/PHONE-BOOT-H/blindzone")
p = d.add_paragraph()
p.add_run("가점 신청: ① AI 학습도구 = XGBoost 회귀(위험지수 학습·재현)  "
          "② AI 분석도구 = SHAP(기여요인 분석)  ③ 데이터 융합(3종 + 119 교차검증). "
          "생성형 AI(Claude Code)는 코드·문서 보조이며 데이터·지표·해석 판단은 사람이 수행.").bold = True

d.add_heading("1. AI 학습도구 — XGBoost 회귀", level=1)
d.add_paragraph(
    "min-max 정규화 후 가중합으로 정의한 잠재 위험지수를, 입력 변수(사고 빈도·사망사고 비율·"
    "응급기관 거리·도착 추정시간·면적)로 학습·재현했다(xgboost.XGBRegressor). 성능 R²=0.90, "
    "MAE=0.0079 — 사고·사망 예측 성능이 아니라 정의식 재현도다. 학습 코드 backend/src/train.py, "
    "모델 backend/models/xgb_risk_model.pkl.")
pic(d, "01_xgboost_feature_importance.png")
caption(d, "[그림1] XGBoost 변수 중요도(gain) — 사고 건수·응급기관 거리가 상위 기여")

d.add_heading("2. AI 분석도구 — SHAP TreeExplainer", level=1)
d.add_paragraph(
    "학습된 모델에 SHAP로 시군구별 기여요인(증가/감소 방향 포함)을 분해해, 시민 모드에서 "
    "'왜 위험한가'로 노출한다. 인제군 예: 응급기관 거리(+0.173) > 사고 빈도(+0.072) > "
    "사망사고 비율(+0.065).")
pic(d, "02_shap_summary.png")
caption(d, "[그림2] SHAP 요약 — 응급기관 거리(ems_distance_km)가 위험지수의 핵심 요인")
pic(d, "03_shap_bar.png")
caption(d, "[그림3] 변수별 평균 기여도(mean |SHAP value|)")
pic(d, "04_shap_inje_waterfall.png")
caption(d, "[그림4] 인제군 위험지수의 변수별 SHAP 분해 — 응급거리 기여 최대")

d.add_heading("3. 데이터 융합 — 3종 융합 + 119 교차검증", level=1)
tbl = d.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 1"
tbl.rows[0].cells[0].text = "데이터(보유기관)"
tbl.rows[0].cells[1].text = "시군구 단위 활용"
fusion = [
    ("TAAS 사고다발지역 (한국도로교통공단)", "다발지점 → 시군구 집계: 사고 건수·사망사고 비율"),
    ("응급의료기관 정보 (국립중앙의료원)", "시군구 중심 → 최근접 응급기관 거리"),
    ("시군구 행정경계 (통계청 센서스경계)", "공간 결합 기준 폴리곤 (CC BY-NC-ND)"),
    ("119 구급통계 (소방청)", "risk_index 변수 미사용. 도착시간 현실성·다발지점 한계 교차검증(인제 1건 vs 119 188건)"),
]
for a, b in fusion:
    c = tbl.add_row().cells
    c[0].text = a
    c[1].text = b

d.add_heading("4. (참고) 생성형 AI 개발보조 — Claude Code", level=1)
d.add_paragraph(
    "코드 작성·오류 수정·배포 설정·문서화 보조에 Claude Code를 활용했다. 데이터 선정, 위험지수·"
    "가중치 설계, 발견 해석, 정직성 표현, 정책적 의미 도출은 사람이 직접 수행했다. AI가 실질적으로 "
    "보조한 주요 커밋에 Co-Authored-By를 명시했으며 자동 commit 없이 검토 후 반영했다.")
tbl2 = d.add_table(rows=1, cols=2)
tbl2.style = "Light Grid Accent 1"
tbl2.rows[0].cells[0].text = "커밋(선별)"
tbl2.rows[0].cells[1].text = "내용"
for h, t in [
    ("08a56a1", "외부검증·도로 실거리 robust 재산출 + 가점표현 정직화"),
    ("84d3845", "고령 수요 교차분석 + 수혜인구"),
    ("68a5a76", "다발지점 vs 전체 교통사고 대조 분석"),
    ("f08ee6b", "정적 fallback + 배포(백엔드 다운 시에도 데모 동작)"),
    ("3ea4a71", "HF Spaces Docker 배포 설정"),
]:
    c = tbl2.add_row().cells
    c[0].text = h
    c[1].text = t
d.add_paragraph(
    "주요 AI 보조 커밋 16건을 선별해 증빙하며, 전체 커밋 이력은 GitHub 저장소 기준으로 확인 "
    "가능하다. (선택) Claude Code 대화 화면·GitHub 커밋 상세(Co-Authored-By) 캡처를 추가하면 보강된다.")

d.save(str(DESK / "BlindZone_AI활용증빙.docx"))
print("saved:", DESK / "BlindZone_AI활용증빙.docx")


# ───────────── 2) 시제품 포트폴리오 ─────────────
d2 = Document()
set_kfont(d2)
d2.add_heading("제품·서비스 개발 기획서 별첨: BlindZone 시제품 포트폴리오", 0)
for line in [
    "BlindZone은 사고다발지도에 잘 드러나지 않는 응급 접근 사각지대를 찾는 웹서비스입니다.",
    "사고 빈도·사망 위험·응급의료 접근성을 결합해 시군구별 잠재 위험지수를 산출합니다.",
    "SHAP 설명과 가상 응급거점 시뮬레이터로 지자체의 우선점검·배치 의사결정을 지원합니다.",
]:
    d2.add_paragraph(line, style="List Bullet")
d2.add_paragraph("라이브: https://blindzone-brown.vercel.app  ·  "
                 "API: https://hananhan-blindzone-backend.hf.space  ·  "
                 "코드: https://github.com/PHONE-BOOT-H/blindzone")

shots = [
    ("10_demo_main.png", "화면 1 — 메인: 전국 위험지도 + TOP10",
     "전국 252개 시군구 잠재 위험지수. 상위는 대부분 대도시(이미 알려진 사고다발지)."),
    ("11_demo_detail.png", "화면 2 — 시군구 상세 + '왜 위험한가'(SHAP)",
     "옹진군 선택: 다발지점 사고 0건이나 응급거리 75.3km. SHAP가 응급 접근성을 핵심 요인으로 "
     "설명. 사각지대 대조·TOP10도 함께 표시."),
    ("12_demo_policy.png", "화면 3 — 정책 시뮬레이터",
     "가상 응급거점 추가 시 거리 재계산 → 위험지수 변화(거리 기반 접근성 민감도). "
     "예: 인제 인근 거점 1곳 추가 시 응급거리 24.2→8.2km·위험지수 약 -58%·48개 시군구 개선 / "
     "옹진 인근 거점 시 75.3→15.4km·약 -81%·47개 시군구 개선."),
    ("13_demo_about.png", "화면 4 — 서비스 소개(About)",
     "무엇을 푸는지·어떻게 찾는지·다섯 가지 검증·한계를 일반 사용자도 이해하게 정리."),
]
for fn, title, desc in shots:
    d2.add_heading(title, level=1)
    d2.add_paragraph(desc)
    pic(d2, fn)

d2.add_heading("기술 구성 · 데이터 · 한계", level=1)
d2.add_paragraph("프론트엔드: Next.js 14 + MapLibre GL + deck.gl (Vercel) / "
                 "백엔드: FastAPI + XGBoost + SHAP (Hugging Face Spaces, Docker). "
                 "백엔드 cold start·중단 시에도 정적 스냅샷으로 핵심 화면이 동작하도록 설계.")
d2.add_paragraph("데이터: TAAS 사고다발(도로교통공단)·응급의료기관(국립중앙의료원)·행정경계(통계청) "
                 "3종 융합 + 전체 교통사고 통계·119 구급통계·고령인구현황으로 검증.")
d2.add_paragraph("한계: 사고·사망을 예측하지 않는 현장검증 우선순위 도구. 응급 접근성은 거리 기반 "
                 "추정(출동·수용·이송체계 미반영, 옹진 도서는 해상이송 미반영).")

d2.save(str(DESK / "BlindZone_시제품포트폴리오.docx"))
print("saved:", DESK / "BlindZone_시제품포트폴리오.docx")
print("완료")
